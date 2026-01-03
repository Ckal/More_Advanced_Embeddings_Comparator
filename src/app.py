import os
import time
import pdfplumber
import docx
import nltk
import gradio as gr
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.embeddings import CohereEmbeddings
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS, Chroma
from langchain_text_splitters import RecursiveCharacterTextSplitter, TokenTextSplitter
from typing import List, Dict, Any
import pandas as pd
import numpy as np
import re
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from nltk.stem import SnowballStemmer
import jellyfish
from gensim.models import Word2Vec
from gensim.models.fasttext import FastText
from collections import Counter
from tokenizers import Tokenizer, models, trainers
from tokenizers.models import WordLevel
from tokenizers.trainers import WordLevelTrainer
from tokenizers.pre_tokenizers import Whitespace
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.manifold import TSNE
from sklearn.metrics import silhouette_score
from scipy.stats import spearmanr
from functools import lru_cache
from langchain.retrievers import MultiQueryRetriever
from langchain_huggingface.llms import HuggingFacePipeline
from transformers import pipeline
from sentence_transformers import SentenceTransformer, util
from sklearn.model_selection import ParameterGrid
from tqdm import tqdm
import random
from huggingface_hub import login
from typing import List, Tuple, Optional


#hf_token = os.getenv("hf_token")
#login(token=hf_token)

# Define the model pipeline with additional generation parameters
#model_pipeline = pipeline(
#   # model="meta-llama/Llama-3.2-1B",
#    model="dunzhang/stella_en_1.5B_v5",
#    #pad_token_id=50256,
#    #use_auth_token=hf_token,
#    #max_length=1000,  # You can increase this if needed
#    max_new_tokens=900  # Limit how many tokens are generated
#)

# Use the pipeline in HuggingFacePipeline
#llm = HuggingFacePipeline(pipeline=model_pipeline)
###################


#llm = HuggingFacePipeline.from_model_id(
#    model_id="bigscience/bloom-1b7",
#    task="text-generation",
#    model_kwargs={"temperature": 0,  "max_length":1200, "do_sample":True},
#)


##### Alternative
from transformers import pipeline
import torch
from transformers import AutoTokenizer, AutoModelForCausalLM, BitsAndBytesConfig

#READER_MODEL_NAME = "HuggingFaceH4/zephyr-7b-beta" #    model="dunzhang/stella_en_1.5B_v5",


#bnb_config = BitsAndBytesConfig(
#    load_in_4bit=True,
#    bnb_4bit_use_double_quant=True,
#    bnb_4bit_quant_type="nf4",
#    bnb_4bit_compute_dtype=torch.bfloat16,
#)
#rmodel = AutoModelForCausalLM.from_pretrained(READER_MODEL_NAME, quantization_config=bnb_config)
#tokenizer = AutoTokenizer.from_pretrained(READER_MODEL_NAME)

#llm = pipeline(
#    model=rmodel,
#    tokenizer=tokenizer,
#    task="text-generation",
#    do_sample=True,
#    temperature=0.2,
#    repetition_penalty=1.1,
#    return_full_text=False,
#    max_new_tokens=500,
#)
#####
from huggingface_hub import InferenceClient

#repo_id = "meta-llama/Llama-3.2-1B-Instruct"

#llm = InferenceClient(model=repo_id, timeout=120)

# Test your LLM client
#llm_client.text_generation(prompt="How are you today?", max_new_tokens=20)

# NLTK Resource Download
def download_nltk_resources():
    resources = ['punkt', 'stopwords', 'snowball_data', 'wordnet']
    for resource in resources:
        try:
            nltk.download(resource, quiet=False)
        except Exception as e:
            print(f"Failed to download {resource}: {str(e)}")

download_nltk_resources()

#nltk.download('punkt')
print("------ Strinage ------ ")
FILES_DIR = './files'


DEFAULT_MODELS = {
    "HuggingFace": [
        "paraphrase-miniLM",
        "paraphrase-mpnet",
        "all-MiniLM-L6-v2"
    ],
    "OpenAI": [
        "text-embedding-ada-002"
    ],
    "Cohere": [
        "embed-multilingual-v2.0"
    ]
}

# Model Management
class ModelManager:
    def __init__(self):
        self.rankings: Dict[str, float] = {}
        self.model_stats: Dict[str, Dict[str, Any]] = {}
        self.models = {
            'HuggingFace': {
                'e5-base-de': "danielheinz/e5-base-sts-en-de",
                'paraphrase-miniLM': "paraphrase-multilingual-MiniLM-L12-v2",
                'paraphrase-mpnet': "paraphrase-multilingual-mpnet-base-v2",
                'gte-large': "gte-large",
                'gbert-base': "gbert-base"
            },
            'OpenAI': {
                'text-embedding-ada-002': "text-embedding-ada-002"
            },
            'Cohere': {
                'embed-multilingual-v2.0': "embed-multilingual-v2.0"
            }
        }


    def update_model_ranking(self, model_id: str, score: float, feedback: str = None):
        """Update model ranking based on performance and optional feedback"""
        current_score = self.rankings.get(model_id, 0.0)
        # Weighted average of current score and new score
        self.rankings[model_id] = 0.7 * current_score + 0.3 * score

        if feedback:
            if model_id not in self.model_stats:
                self.model_stats[model_id] = {"feedback_count": 0, "feedback": []}
            self.model_stats[model_id]["feedback_count"] += 1
            self.model_stats[model_id]["feedback"].append(feedback)

    def get_top_models(self, n: int = 5) -> List[Tuple[str, float]]:
        """Get top n ranked models"""
        return sorted(self.rankings.items(), key=lambda x: x[1], reverse=True)[:n]

    def get_model_stats(self, model_id: str) -> Dict[str, Any]:
        """Get statistics for a specific model"""
        return self.model_stats.get(model_id, {})


    def add_model(self, provider, name, model_path):
        if provider not in self.models:
            self.models[provider] = {}
        self.models[provider][name] = model_path

    def remove_model(self, provider, name):
        if provider in self.models and name in self.models[provider]:
            del self.models[provider][name]

    def get_model(self, provider, name):
        return self.models.get(provider, {}).get(name)

    def list_models(self):
        return {provider: list(models.keys()) for provider, models in self.models.items()}

model_manager = ModelManager()

# File Handling
import os
import json
import csv
import xml.etree.ElementTree as ET
import openpyxl  # for handling .xlsx files
import pdfplumber
import docx

# File Handling
class FileHandler:
    @staticmethod
    def extract_text(file_path):
        ext = os.path.splitext(file_path)[-1].lower()
        if ext == '.pdf':
            return FileHandler._extract_from_pdf(file_path)
        elif ext == '.docx':
            return FileHandler._extract_from_docx(file_path)
        elif ext == '.txt':
            return FileHandler._extract_from_txt(file_path)
        elif ext == '.xml':
            return FileHandler._extract_from_xml(file_path)
        elif ext == '.json':
            return FileHandler._extract_from_json(file_path)
        elif ext == '.xlsx':
            return FileHandler._extract_from_xlsx(file_path)
        elif ext == '.csv':
            return FileHandler._extract_from_csv(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    @staticmethod
    def _extract_from_pdf(file_path):
        with pdfplumber.open(file_path) as pdf:
            return ' '.join([page.extract_text() for page in pdf.pages])

    @staticmethod
    def _extract_from_docx(file_path):
        doc = docx.Document(file_path)
        return ' '.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _extract_from_txt(file_path):
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    @staticmethod
    def _extract_from_xml(file_path):
        tree = ET.parse(file_path)
        root = tree.getroot()
        return FileHandler._extract_xml_text(root)

    @staticmethod
    def _extract_xml_text(element):
        # Recursively extract text from XML elements
        text = element.text or ""
        for child in element:
            text += FileHandler._extract_xml_text(child)
        return text.strip()

    @staticmethod
    def _extract_from_json(file_path):
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            return json.dumps(data, indent=4)  # Pretty print JSON for readability

    @staticmethod
    def _extract_from_xlsx(file_path):
        workbook = openpyxl.load_workbook(file_path)
        sheet = workbook.active
        data = []
        for row in sheet.iter_rows(values_only=True):
            data.append('\t'.join([str(cell) for cell in row if cell is not None]))
        return '\n'.join(data)

    @staticmethod
    def _extract_from_csv(file_path):
        with open(file_path, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            data = []
            for row in reader:
                data.append(','.join(row))
            return '\n'.join(data)



# Text Processing
def simple_tokenize(text):
    return text.split()

def preprocess_text(text, lang='german', apply_preprocessing=False):
    if not apply_preprocessing:
        return text

    text = text.lower()
    text = re.sub(r'[^a-zA-Z\s]', '', text)

    try:
        tokens = word_tokenize(text, language=lang)
    except LookupError:
        print(f"Warning: NLTK punkt tokenizer for {lang} not found. Using simple tokenization.")
        tokens = simple_tokenize(text)

    try:
        stop_words = set(stopwords.words(lang))
    except LookupError:
        print(f"Warning: Stopwords for {lang} not found. Skipping stopword removal.")
        stop_words = set()
    tokens = [token for token in tokens if token not in stop_words]

    try:
        stemmer = SnowballStemmer(lang)
        tokens = [stemmer.stem(token) for token in tokens]
    except ValueError:
        print(f"Warning: SnowballStemmer for {lang} not available. Skipping stemming.")

    return ' '.join(tokens)

def phonetic_match(text, query, method='levenshtein_distance', apply_phonetic=False):
    if not apply_phonetic:
        return 0
    if method == 'levenshtein_distance':
        text_phonetic = jellyfish.soundex(text)
        query_phonetic = jellyfish.soundex(query)
        return jellyfish.levenshtein_distance(text_phonetic, query_phonetic)
    return 0


from typing import List, Union
import torch
from transformers import AutoTokenizer
import numpy as np
from nltk.tokenize import word_tokenize
from nltk.corpus import wordnet
import nltk

def optimize_query(
    query: str,
    query_optimization_model: str,  # Added to match your signature = "google/flan-t5-small"
    chunks: List[str],
    embedding_model: str,
    vector_store_type: str,  # Added to match your signature
    search_type: str,  # Added to match your signature
    top_k: int = 3,
    use_gpu: bool = False
) -> str:
    """
    CPU-optimized version of query expansion using a small language model.

    Args:
        query: Original search query
        query_optimization_model: Name or path of the model to use for optimization
        chunks: List of text chunks to search through
        embedding_model: Name of the embedding model being used
        vector_store_type: Type of vector store being used
        search_type: Type of search being performed
        top_k: Number of expansion terms to add
        use_gpu: Whether to use GPU if available (defaults to False for CPU)

    Returns:
        Expanded query string
    """
    try:
        # Set device
        device = "cuda" if use_gpu and torch.cuda.is_available() else "cpu"

        # 1. Basic text preprocessing (CPU-based)
        tokens = word_tokenize(query.lower())

        # 2. WordNet synonyms expansion (CPU-based)
        expanded_terms = set()
        for token in tokens:
            # Limit synonym lookup to save CPU resources
            synsets = wordnet.synsets(token)[:1]  # Take only top synset per word
            for syn in synsets:
                # Limit number of lemmas
                expanded_terms.update([lemma.name() for lemma in syn.lemmas()[:2]])

        # 3. Use provided model with reduced complexity
            try:
                # Initialize the pipeline with the chosen model
                llm_pipeline = pipeline(model="meta-llama/Llama-3.2-1B-Instruct", device='cpu')

                # Define prompt for the assistant, making it context-specific
                prompt = f'''
                <|start_header_id|>system<|end_header_id|>
                You are an expert in enhancing user input for vector store retrieval.
                Enhance the followinf search query with relevant terms.

                show me just the new term. You SHOULD NOT include any other text in the response.

                <|eot_id|><|start_header_id|>user<|end_header_id|>
                {query}
                <|eot_id|><|start_header_id|>assistant<|end_header_id|>
                '''

                # Get suggested settings from the LLM
                suggested_settings = llm_pipeline(
                    prompt,
                    do_sample=True,
                    top_k=10,
                    num_return_sequences=1,
                    return_full_text=False,
                    max_new_tokens=1900,    # Control the length of the output
                    truncation=True  # Enable truncation
                )

                # Extract the settings from the generated response
                generated_text = suggested_settings[0].get('generated_text', '')
                print(generated_text)  # For debugging, ensure text output is as expected

            except Exception as model_error:
                print(f"LLM-based expansion failed: {str(model_error)}")
                generated_text = "Default settings could not be generated."  # Fallback message or settings


            # 4. Combine original and expanded terms
            final_terms = set(tokens)
            final_terms.update(expanded_terms)
            if generated_text != query:
                final_terms.update(word_tokenize(generated_text.lower()))

        # 5. Remove stopwords and select top_k most relevant terms
        stopwords = set(['the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to'])
        final_terms = [term for term in final_terms if term not in stopwords]

        # Combine with original query
        generated_text = f"{query} {' '.join(list(final_terms)[:top_k])}"
        print(generated_text)
        # Clean up
       # llm_pipeline = None

        return generated_text.strip() #[Document(page_content=generated_text.strip())]

    except Exception as e:
        print(f"Query optimization failed: {str(e)}")
        return query #[Document(page_content=query)]  # Return original query if optimization fails



# Example usage
"""
chunks = ["sample text chunk 1", "sample text chunk 2"]
query = "machine learning algorithms"
optimized_query = optimize_query(
    query=query,
    chunks=chunks,
    embedding_model="sentence-transformers/all-MiniLM-L6-v2",
    use_gpu=False  # Explicitly use CPU
)
"""


def create_custom_embedding(texts, model_type='word2vec', vector_size=100, window=5, min_count=1):
    tokenized_texts = [text.split() for text in texts]

    if model_type == 'word2vec':
        model = Word2Vec(sentences=tokenized_texts, vector_size=vector_size, window=window, min_count=min_count, workers=4)
    elif model_type == 'fasttext':
        model = FastText(sentences=tokenized_texts, vector_size=vector_size, window=window, min_count=min_count, workers=4)
    else:
        raise ValueError("Unsupported model type")

    return model

class CustomEmbeddings(HuggingFaceEmbeddings):
    def __init__(self, model_path):
        self.model = Word2Vec.load(model_path)  # or FastText.load() for FastText models

    def embed_documents(self, texts):
        return [self.model.wv[text.split()] for text in texts]

    def embed_query(self, text):
        return self.model.wv[text.split()]

# Custom Tokenizer
def create_custom_tokenizer(file_path, model_type='WordLevel', vocab_size=10000, special_tokens=None):
    with open(file_path, 'r', encoding='utf-8') as f:
        text = f.read()

    if model_type == 'WordLevel':
        tokenizer = Tokenizer(WordLevel(unk_token="[UNK]"))
    elif model_type == 'BPE':
        tokenizer = Tokenizer(models.BPE(unk_token="[UNK]"))
    elif model_type == 'Unigram':
        tokenizer = Tokenizer(models.Unigram())
    else:
        raise ValueError(f"Unsupported tokenizer model: {model_type}")

    tokenizer.pre_tokenizer = Whitespace()

    special_tokens = special_tokens or ["[UNK]", "[CLS]", "[SEP]", "[PAD]", "[MASK]"]
    trainer = trainers.WordLevelTrainer(special_tokens=special_tokens, vocab_size=vocab_size)
    tokenizer.train_from_iterator([text], trainer)

    return tokenizer

def custom_tokenize(text, tokenizer):
    return tokenizer.encode(text).tokens

# Embedding and Vector Store
#@lru_cache(maxsize=None)

# Helper functions

def get_text_splitter(split_strategy, chunk_size, overlap_size, custom_separators=None):
    if split_strategy == 'token':
        return TokenTextSplitter(chunk_size=chunk_size, chunk_overlap=overlap_size)
    elif split_strategy == 'recursive':
        return RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=overlap_size,
            add_start_index=True,  # If `True`, includes chunk's start index in metadata
            strip_whitespace=True,  # If `True`, strips whitespace from the start and end of every document
            separators=custom_separators or ["\n\n", "\n", " ", ""]
        )
    else:
        raise ValueError(f"Unsupported split strategy: {split_strategy}")

def get_embedding_model(model_type, model_name):
    model_path = model_manager.get_model(model_type, model_name)
    if model_type == 'HuggingFace':
        return HuggingFaceEmbeddings(
            model_name=model_path,
            multi_process=True,
           # model_kwargs={"device": "cpu"},
            #encode_kwargs={"normalize_embeddings": True},  # Set `True` for cosine similarity
        )
    elif model_type == 'OpenAI':
        return OpenAIEmbeddings(model=model_path)
    elif model_type == 'Cohere':
        return CohereEmbeddings(model=model_path)
    else:
        raise ValueError(f"Unsupported model type: {model_type}")

def get_vector_store(vector_store_type, chunks, embedding_model):
    chunks_tuple = tuple(chunks)
    if vector_store_type == 'FAISS':
        return FAISS.from_texts(chunks, embedding_model)
    elif vector_store_type == 'Chroma':
        return Chroma.from_texts(chunks, embedding_model)
    else:
        raise ValueError(f"Unsupported vector store type: {vector_store_type}")

def get_retriever(vector_store, search_type, search_kwargs):
    if search_type == 'similarity':
        return vector_store.as_retriever(search_type="similarity", search_kwargs=search_kwargs)
    elif search_type == 'mmr':
        return vector_store.as_retriever(search_type="mmr", search_kwargs=search_kwargs)
    elif search_type == 'custom':
        return vector_store.as_retriever(search_type="similarity", search_kwargs=search_kwargs)
    else:
        raise ValueError(f"Unsupported search type: {search_type}")

def custom_similarity(query_embedding, doc_embedding, query, doc_text, phonetic_weight=0.3):
    embedding_sim = np.dot(query_embedding, doc_embedding) / (np.linalg.norm(query_embedding) * np.linalg.norm(doc_embedding))
    phonetic_sim = phonetic_match(doc_text, query)
    combined_sim = (1 - phonetic_weight) * embedding_sim + phonetic_weight * phonetic_sim
    return combined_sim

def _create_vector_store(vector_store_type, chunks_tuple, embedding_model):
    chunks = list(chunks_tuple)

    if vector_store_type == 'FAISS':
        return FAISS.from_texts(chunks, embedding_model)
    elif vector_store_type == 'Chroma':
        return Chroma.from_texts(chunks, embedding_model)
    else:
        raise ValueError(f"Unsupported vector store type: {vector_store_type}")


# Main Processing Functions
def process_files(file_path, model_type, model_name, split_strategy, chunk_size, overlap_size, custom_separators, lang='german', apply_preprocessing=False, custom_tokenizer_file=None, custom_tokenizer_model=None, custom_tokenizer_vocab_size=10000, custom_tokenizer_special_tokens=None):
    if file_path:
        text = FileHandler.extract_text(file_path)
    else:
        text = ""
        for file in os.listdir(FILES_DIR):
            file_path = os.path.join(FILES_DIR, file)
            text += FileHandler.extract_text(file_path)

    if custom_tokenizer_file:
        tokenizer = create_custom_tokenizer(custom_tokenizer_file, custom_tokenizer_model, custom_tokenizer_vocab_size, custom_tokenizer_special_tokens)
        text = ' '.join(custom_tokenize(text, tokenizer))
    elif apply_preprocessing:
        text = preprocess_text(text, lang)

    text_splitter = get_text_splitter(split_strategy, chunk_size, overlap_size, custom_separators)
    chunks = text_splitter.split_text(text)

    embedding_model = get_embedding_model(model_type, model_name)

    return chunks, embedding_model, len(text.split())

def search_embeddings(chunks, embedding_model, vector_store_type, search_type, query, top_k, expected_result=None, lang='german', apply_phonetic=False, phonetic_weight=0.3):
    preprocessed_query = preprocess_text(query, lang) if apply_phonetic else query

    vector_store = get_vector_store(vector_store_type, chunks, embedding_model)
    retriever = get_retriever(vector_store, search_type, {"k": top_k})

    start_time = time.time()
    results = retriever.invoke(preprocessed_query)

    #this should be optional
    def score_result(doc):
        base_score = vector_store.similarity_search_with_score(doc.page_content, k=1)[0][1]

        # Add bonus for containing expected result
        expected_bonus = 0.3 if expected_result and expected_result in doc.page_content else 0

        if apply_phonetic:
            phonetic_score = phonetic_match(doc.page_content, query)
            return (1 - phonetic_weight) * base_score + phonetic_weight * phonetic_score + expected_bonus
        else:
            return base_score + expected_bonus

    results = sorted(results, key=score_result, reverse=True)
    end_time = time.time()

    embeddings = []
    for doc in results:
        if hasattr(doc, 'embedding'):
            embeddings.append(doc.embedding)
        else:
            embeddings.append(None)

    results_df = pd.DataFrame({
        'content': [doc.page_content for doc in results],
        'embedding': embeddings,
        'length': [len(doc.page_content) for doc in results],
        'contains_expected': [expected_result in doc.page_content if expected_result else None for doc in results]
    })

    return results_df, end_time - start_time, vector_store, results

# Enhanced Result Analysis
class ResultAnalyzer:
    @staticmethod
    def calculate_statistics(results, search_time, vector_store, num_tokens, embedding_model, query,
                           top_k, expected_result=None, model_feedback=None):
        stats = {
            "num_results": len(results),
            "avg_content_length": np.mean([len(doc.page_content) for doc in results]) if results else 0,
            "min_content_length": min([len(doc.page_content) for doc in results]) if results else 0,
            "max_content_length": max([len(doc.page_content) for doc in results]) if results else 0,
            "search_time": search_time,
            "num_tokens": num_tokens,
            "embedding_dimension": len(embedding_model.embed_query(query)),
            "top_k": top_k,
        }

        # Add vector store statistics
        try:
            if hasattr(vector_store, '_index'):
                stats["vector_store_size"] = vector_store._index.ntotal
            elif hasattr(vector_store, '_collection'):
                stats["vector_store_size"] = len(vector_store._collection.get())
        except:
            stats["vector_store_size"] = "N/A"

        # Add expected result statistics if provided
        if expected_result:
            stats["contains_expected"] = any(expected_result in doc.page_content for doc in results)
            stats["expected_result_rank"] = next((i for i, doc in enumerate(results)
                                                if expected_result in doc.page_content), -1) + 1

        # Calculate diversity metrics for larger result sets
        if len(results) > 3:  # Changed from 1000 to make it more practical
            embeddings = [embedding_model.embed_query(doc.page_content) for doc in results]
            stats["result_diversity"] = ResultAnalyzer._calculate_diversity(embeddings)
            stats["silhouette_score"] = ResultAnalyzer._calculate_silhouette(embeddings)
        else:
            stats["result_diversity"] = "N/A"
            stats["silhouette_score"] = "N/A"

        # Add ranking correlation
        query_embedding = embedding_model.embed_query(query)
        result_embeddings = [embedding_model.embed_query(doc.page_content) for doc in results]
        similarities = [np.inner(query_embedding, emb) for emb in result_embeddings]
        if len(similarities) > 1:
            rank_correlation, _ = spearmanr(similarities, range(len(similarities)))
            stats["rank_correlation"] = rank_correlation
        else:
            stats["rank_correlation"] = "N/A"

        # Add model feedback if provided
        if model_feedback:
            stats["model_feedback"] = model_feedback

        return stats

    @staticmethod
    def _calculate_diversity(embeddings: List[np.ndarray]) -> float:
        """Calculate diversity score for embeddings"""
        embeddings_array = np.array(embeddings)
        pairwise_similarities = np.inner(embeddings_array, embeddings_array)
        return 1 - np.mean(pairwise_similarities[np.triu_indices(len(embeddings), k=1)])

    @staticmethod
    def _calculate_silhouette(embeddings: List[np.ndarray]) -> float:
        """Calculate silhouette score for embeddings"""
        if len(embeddings) < 3:
            return 0.0
        try:
            return silhouette_score(embeddings, range(len(embeddings)))
        except:
            return 0.0





# Visualization
def visualize_results(results_df, stats_df):
    # Add model column if not present
    if 'model' not in stats_df.columns:
        stats_df['model'] = stats_df['model_type'] + ' - ' + stats_df['model_name']

    fig, axs = plt.subplots(2, 2, figsize=(20, 20))

    # Handle empty dataframe case
    if len(stats_df) == 0:
        return fig

    # Create plots with error handling
    try:
        sns.barplot(data=stats_df, x='model', y='search_time', ax=axs[0, 0])
        axs[0, 0].set_title('Search Time by Model')
        axs[0, 0].tick_params(axis='x', rotation=45)
    except Exception as e:
        print(f"Error in search time plot: {e}")

    try:
        sns.scatterplot(data=stats_df, x='result_diversity', y='rank_correlation',
                       hue='model', ax=axs[0, 1])
        axs[0, 1].set_title('Result Diversity vs. Rank Correlation')
    except Exception as e:
        print(f"Error in diversity plot: {e}")

    try:
        sns.boxplot(data=stats_df, x='model', y='avg_content_length', ax=axs[1, 0])
        axs[1, 0].set_title('Distribution of Result Content Lengths')
        axs[1, 0].tick_params(axis='x', rotation=45)
    except Exception as e:
        print(f"Error in content length plot: {e}")

    try:
        valid_embeddings = results_df['embedding'].dropna().values
        if len(valid_embeddings) > 1:
            tsne = TSNE(n_components=2, random_state=42)
            embeddings_2d = tsne.fit_transform(np.vstack(valid_embeddings))
            sns.scatterplot(x=embeddings_2d[:, 0], y=embeddings_2d[:, 1],
                          hue=results_df['Model'][:len(valid_embeddings)],
                          ax=axs[1, 1])
            axs[1, 1].set_title('t-SNE Visualization of Result Embeddings')
        else:
            axs[1, 1].text(0.5, 0.5, "Not enough embeddings for visualization",
                          ha='center', va='center')
    except Exception as e:
        print(f"Error in embedding visualization: {e}")

    plt.tight_layout()
    return fig


#tokenizer = AutoTokenizer.from_pretrained(EMBEDDING_MODEL_NAME)
#lengths = [len(tokenizer.encode(doc.page_content)) for doc in tqdm(docs_processed)]
#fig = pd.Series(lengths).hist()
#plt.title("Distribution of document lengths in the knowledge base (in count of tokens)")
#plt.show()


def optimize_vocabulary(texts, vocab_size=10000, min_frequency=2):
    tokenizer = Tokenizer(models.BPE(unk_token="[UNK]"))

    word_freq = Counter(word for text in texts for word in text.split())

    optimized_texts = [
        ' '.join(word for word in text.split() if word_freq[word] >= min_frequency)
        for text in texts
    ]

    trainer = trainers.BpeTrainer(vocab_size=vocab_size, special_tokens=["[UNK]", "[CLS]", "[SEP]", "[PAD]", "[MASK]"])
    tokenizer.train_from_iterator(optimized_texts, trainer)

    return tokenizer, optimized_texts

import numpy as np
from transformers import TextClassificationPipeline
from typing import List, Union, Any



model = SentenceTransformer('sentence-transformers/all-MiniLM-L6-v2')


def rerank_results(
    results: List[Any],
    query: str,
    reranker: Union[TextClassificationPipeline, Any]
) -> List[Any]:
    """

    """
    if not results:
        return results

    # Step 1: Encode the query and documents using SentenceTransformer
    query_embedding = model.encode(query, convert_to_tensor=True)
    doc_contents = [doc.page_content for doc in results]  # Assuming each result has a `page_content` attribute
    doc_embeddings = model.encode(doc_contents, convert_to_tensor=True)

    # Step 2: Compute cosine similarities between query and document embeddings
    cosine_scores = util.cos_sim(query_embedding, doc_embeddings)[0]  # Shape: (number of documents,)

    # Step 3: Sort documents by similarity score in descending order
    reranked_idx = np.argsort(cosine_scores.cpu().numpy())[::-1]

    # Step 4: Return the reranked documents
    reranked_results = [results[i] for i in reranked_idx]

    return reranked_results


# Main Comparison Function
def compare_embeddings(file, query, embedding_models, custom_embedding_model, split_strategy, chunk_size, overlap_size, custom_separators, vector_store_type, search_type, top_k, expected_result=None, lang='german', apply_preprocessing=True, optimize_vocab=False, apply_phonetic=True, phonetic_weight=0.3, custom_tokenizer_file=None, custom_tokenizer_model=None, custom_tokenizer_vocab_size=10000, custom_tokenizer_special_tokens=None, use_query_optimization=False, query_optimization_model="google/flan-t5-base", use_reranking=False):
    all_results = []
    all_stats = []
    settings = {
        "split_strategy": split_strategy,
        "chunk_size": chunk_size,
        "overlap_size": overlap_size,
        "custom_separators": custom_separators,
        "vector_store_type": vector_store_type,
        "search_type": search_type,
        "top_k": top_k,
        "lang": lang,
        "apply_preprocessing": apply_preprocessing,
        "optimize_vocab": optimize_vocab,
        "apply_phonetic": apply_phonetic,
        "phonetic_weight": phonetic_weight,
        "use_query_optimization": use_query_optimization,
        "query_optimization_model": query_optimization_model,
        "use_reranking": use_reranking
    }

    # Parse the embedding models from the checkbox group
    models = [model.split(':') for model in embedding_models]
    if custom_embedding_model:
        models.append(custom_embedding_model.strip().split(':'))

    for model_type, model_name in models:
        chunks, embedding_model, num_tokens = process_files(
            file.name if file else None,
            model_type,
            model_name,
            split_strategy,
            chunk_size,
            overlap_size,
            custom_separators.split(',') if custom_separators else None,
            lang,
            apply_preprocessing,
            custom_tokenizer_file,
            custom_tokenizer_model,
            int(custom_tokenizer_vocab_size),
            custom_tokenizer_special_tokens.split(',') if custom_tokenizer_special_tokens else None
        )

        if optimize_vocab:
            tokenizer, optimized_chunks = optimize_vocabulary(chunks)
            chunks = optimized_chunks

        search_query = query

        if use_query_optimization:
            optimized_queries = optimize_query(query, query_optimization_model, chunks, embedding_model, vector_store_type, search_type, top_k)
            #query = " ".join(optimized_queries)
            search_query = optimized_queries # " ".join([doc.page_content for doc in optimized_queries])  # Extract text from Document objects

        results, search_time, vector_store, results_raw = search_embeddings(
            chunks,
            embedding_model,
            vector_store_type,
            search_type,
            search_query,
            top_k,
            expected_result,
            lang,
            apply_phonetic,
            phonetic_weight
        )

        if use_reranking:
            reranker = pipeline("text-classification", model="cross-encoder/ms-marco-MiniLM-L-12-v2")
            results_raw = rerank_results(results_raw, query, reranker)

        result_embeddings = [doc.metadata.get('embedding', None) for doc in results_raw]

        stats = ResultAnalyzer.calculate_statistics(results_raw, search_time, vector_store, num_tokens, embedding_model, query, top_k, expected_result)
        stats["model"] = f"{model_type} - {model_name}"
        stats["model_type"] = model_type
        stats["model_name"] = model_name
        stats.update(settings)

        formatted_results = format_results(results_raw, stats)
        for i, result in enumerate(formatted_results):
            result['embedding'] = result_embeddings[i]
            result['length'] = len(result['Content'])
            result['contains_expected'] = expected_result in result['Content'] if expected_result else None

        all_results.extend(formatted_results)
        all_stats.append(stats)

    results_df = pd.DataFrame(all_results)
    stats_df = pd.DataFrame(all_stats)

    fig = visualize_results(results_df, stats_df)

    best_results = analyze_results(stats_df)

    return results_df, stats_df, fig, best_results

def format_results(results, stats):
    formatted_results = []
    for doc in results:
        result = {
            "Model": stats["model"],
            "Content": doc.page_content,
            "Embedding": doc.embedding if hasattr(doc, 'embedding') else None,
            **doc.metadata,
            **{k: v for k, v in stats.items() if k not in ["model"]}
        }
        formatted_results.append(result)
    return formatted_results


#####
from sklearn.model_selection import ParameterGrid
from tqdm import tqdm

# ... (previous code remains the same)

# function for automated testing
def automated_testing(file, query, test_params, expected_result=None):
    all_results = []
    all_stats = []

    param_grid = ParameterGrid(test_params)
    print(param_grid)
    for params in tqdm(param_grid, desc="Running tests"):
        chunks, embedding_model, num_tokens = process_files(
            file.name if file else None,
            params['model_type'],
            params['model_name'],
            params['split_strategy'],
            params['chunk_size'],
            params['overlap_size'],
            params.get('custom_separators', None),
            params['lang'],
            params['apply_preprocessing'],
            params.get('custom_tokenizer_file', None),
            params.get('custom_tokenizer_model', None),
            params.get('custom_tokenizer_vocab_size', 10000),
            params.get('custom_tokenizer_special_tokens', None)
        )

        if params['optimize_vocab']:
            tokenizer, optimized_chunks = optimize_vocabulary(chunks)
            chunks = optimized_chunks

        if params['use_query_optimization']:
            optimized_queries = optimize_query(query, params['query_optimization_model'], chunks ,  embedding_model ,  params['vector_store_type'] , params['search_type'] ,  params['top_k'] )

            #optimized_queries = optimize_query(query, )
            query = " ".join(optimized_queries)

        results, search_time, vector_store, results_raw = search_embeddings(
            chunks,
            embedding_model,
            params['vector_store_type'],
            params['search_type'],
            query,
            params['top_k'],
            expected_result,
            params['lang'],
            params['apply_phonetic'],
            params['phonetic_weight']
        )

        if params['use_reranking']:
            reranker = pipeline("text-classification", model="cross-encoder/ms-marco-MiniLM-L-12-v2")
            results_raw = rerank_results(results_raw, query, reranker)

        stats = ResultAnalyzer.calculate_statistics(results_raw, search_time, vector_store, num_tokens, embedding_model, query, params['top_k'], expected_result)
        stats["model"] = f"{params['model_type']} - {params['model_name']}"
        stats["model_type"] = params['model_type']
        stats["model_name"] = params['model_name']
        stats.update(params)

        all_results.extend(format_results(results_raw, stats))
        all_stats.append(stats)

    return pd.DataFrame(all_results), pd.DataFrame(all_stats)


# Function to analyze results and propose best model and settings
def analyze_results(stats_df):
    metric_weights = {
        'search_time': -0.3,
        'result_diversity': 0.2,
        'rank_correlation': 0.3,
        'silhouette_score': 0.2,
        'contains_expected': 0.5,  # High weight for containing the expected result
        'expected_result_rank': -0.4  # Lower rank (closer to 1) is better
    }
    if stats_df.empty:
      print("stats_df is empty. Cannot compute best configuration.")
      return None

    for metric in metric_weights.keys():

        if metric in stats_df.columns:
            stats_df[metric] = pd.to_numeric(stats_df[metric], errors='coerce')
        else:
            stats_df[metric] = 0
            print("Column 'search_time' is missing in stats_df.")



    stats_df['weighted_score'] = sum(
        stats_df[metric].fillna(0) * weight
        for metric, weight in metric_weights.items()
    )

    best_config = stats_df.loc[stats_df['weighted_score'].idxmax()]

    recommendations = {
        'best_model': f"{best_config['model_type']} - {best_config['model_name']}",
        'best_settings': {
            'split_strategy': best_config['split_strategy'],
            'chunk_size': int(best_config['chunk_size']),
            'overlap_size': int(best_config['overlap_size']),
            'vector_store_type': best_config['vector_store_type'],
            'search_type': best_config['search_type'],
            'top_k': int(best_config['top_k']),
            'optimize_vocab': bool(best_config['optimize_vocab']),
            'use_query_optimization': bool(best_config['use_query_optimization']),
            'use_reranking': bool(best_config['use_reranking']),
            'lang': best_config['lang'],
            'apply_preprocessing': bool(best_config['apply_preprocessing']),
            'apply_phonetic': bool(best_config['apply_phonetic']),
            'phonetic_weight': float(best_config['phonetic_weight'])
        },
        'performance_summary': {
            'search_time': float(best_config['search_time']),
            'result_diversity': float(best_config['result_diversity']),
            'rank_correlation': float(best_config['rank_correlation']),
            'silhouette_score': float(best_config['silhouette_score']),
            'contains_expected': bool(best_config['contains_expected']),
            'expected_result_rank': int(best_config['expected_result_rank'])
        }
    }

    return recommendations

    ####
import ast

def get_llm_suggested_settings(file, num_chunks=1):
    if not file:
        return {"error": "No file uploaded"}

    chunks, _, _ = process_files(
        file.name,
        'HuggingFace',
        'paraphrase-miniLM',
        'recursive',
        250,
        50,
        custom_separators=None
    )

    # Select a few random chunks
    sample_chunks = random.sample(chunks, min(num_chunks, len(chunks)))


    llm_pipeline = pipeline(model="meta-llama/Llama-3.2-1B-Instruct", device='cpu')


    prompt=f'''
    <|start_header_id|>system<|end_header_id|>
    You are an expert in information retrieval.
    You know about strenghs and weaknesses of all models.

    Given the following text chunks from a document,
    suggest optimal settings for an embedding-based search system. The settings should include:

    1. Embedding model type and name
    2. Split strategy (token or recursive)
    3. Chunk size
    4. Overlap size
    5. Vector store type (FAISS or Chroma)
    6. Search type (similarity, mmr, or custom)
    7. Top K results to retrieve
    8. Whether to apply preprocessing
    9. Whether to optimize vocabulary
    10. Whether to apply phonetic matching

    Expected output format:
    {{
        "embedding_models": "embedding_model_type:embedding_model_name",
        "split_strategy": "token or recursive",
        "chunk_size": 250,
        "overlap_size": 50,
        "vector_store_type": "FAISS or Chroma",
        "search_type": "similarity, mmr, or custom",
        "top_k": 5,
        "apply_preprocessing": True,
        "optimize_vocab": True,
        "apply_phonetic": False,
        "phonetic_weight": 0.3  #
    }}

    Provide your suggestions in a Python dictionary format.

    show me settings You SHOULD NOT include any other text in the response.
    Fill out the seeting and chose usefull values.
    Respect the users use cases and content snipet. Choose the setting based on the chunks

    <|eot_id|><|start_header_id|>user<|end_header_id|>
    User user case:
    {"small local", "large total context", ...}

    total content lenght:
    {len(' '.join(chunks))}

    Content snipet:
    {' '.join(sample_chunks)}
    <|eot_id|><|start_header_id|>assistant<|end_header_id|>
    '''
    suggested_settings = llm_pipeline(
            prompt,
            do_sample=True,
            top_k=10,
            num_return_sequences=1,
            return_full_text=False,
            max_new_tokens=1900,    # Control the length of the output,
            truncation=True,  # Enable truncation
        )

    print(suggested_settings[0]['generated_text'])
    # Safely parse the generated text to extract the dictionary
    try:
        # Using ast.literal_eval for safe parsing
        settings_dict = ast.literal_eval(suggested_settings[0]['generated_text'])

        # Convert the settings to match the interface inputs
        return {
            "embedding_models": settings_dict["embedding_models"],
            "split_strategy": settings_dict["split_strategy"],
            "chunk_size": settings_dict["chunk_size"],
            "overlap_size": settings_dict["overlap_size"],
            "vector_store_type": settings_dict["vector_store_type"],
            "search_type": settings_dict["search_type"],
            "top_k": settings_dict["top_k"],
            "apply_preprocessing": settings_dict["apply_preprocessing"],
            "optimize_vocab": settings_dict["optimize_vocab"],
            "apply_phonetic": settings_dict["apply_phonetic"],
            "phonetic_weight": settings_dict.get("phonetic_weight", 0.3)  # Set default if not provided
        }
    except Exception as e:
        print(f"Error parsing LLM suggestions: {e}")
        return {"error": "Failed to parse LLM suggestions"}


def update_inputs_with_llm_suggestions(suggestions):
    if suggestions is None or "error" in suggestions:
        return [gr.update() for _ in range(11)]  # Return no updates if there's an error or None

    return [
        gr.update(value=[suggestions["embedding_models"]]),  # embedding_models_input
        gr.update(value=suggestions["split_strategy"]),      # split_strategy_input
        gr.update(value=suggestions["chunk_size"]),          # chunk_size_input
        gr.update(value=suggestions["overlap_size"]),        # overlap_size_input
        gr.update(value=suggestions["vector_store_type"]),   # vector_store_type_input
        gr.update(value=suggestions["search_type"]),         # search_type_input
        gr.update(value=suggestions["top_k"]),               # top_k_input
        gr.update(value=suggestions["apply_preprocessing"]), # apply_preprocessing_input
        gr.update(value=suggestions["optimize_vocab"]),      # optimize_vocab_input
        gr.update(value=suggestions["apply_phonetic"]),      # apply_phonetic_input
        gr.update(value=suggestions["phonetic_weight"])      # phonetic_weight_input
    ]

def parse_model_selections(default_models, custom_models):
    """
    Parse selected default models and custom models into model configurations

    Args:
        default_models (List[str]): Selected default models in format "type:name"
        custom_models (str): Custom models string with one model per line in format "type:name"

    Returns:
        List[Dict[str, str]]: List of model configurations with 'type' and 'name' keys
    """
    model_configs = []

    # Process default models
    if default_models:
        for model in default_models:
            model_type, model_name = model.split(':')
            model_configs.append({
                'type': model_type,
                'name': model_name
            })

    # Process custom models
    if custom_models:
        custom_model_lines = custom_models.strip().split('\n')
        for line in custom_model_lines:
            if line.strip() and ':' in line:
                model_type, model_name = line.strip().split(':')
                model_configs.append({
                    'type': model_type.strip(),
                    'name': model_name.strip()
                })

    return model_configs

def parse_comma_separated(text):
    """Parse comma-separated values into a list"""
    if not text:
        return []
    return [x.strip() for x in text.split(',') if x.strip()]



# Gradio Interface
def launch_interface(debug=True):
    with gr.Blocks() as iface:
        gr.Markdown("# Advanced Embedding Comparison Tool")

        with gr.Tab("Simple"):
            file_input = gr.File(label="Upload File (Optional)")
            query_input = gr.Textbox(label="Search Query")
            expected_result_input = gr.Textbox(label="Expected Result (Optional)")
            embedding_models_input = gr.CheckboxGroup(
                choices=[
                    "HuggingFace:paraphrase-miniLM",
                    "HuggingFace:paraphrase-mpnet",
                    "OpenAI:text-embedding-ada-002",
                    "Cohere:embed-multilingual-v2.0"
                ],
                label="Embedding Models"
            )
            top_k_input = gr.Slider(1, 10, step=1, value=5, label="Top K")

        with gr.Tab("Advanced"):
            custom_embedding_model_input = gr.Textbox(label="Custom Embedding Model (optional, format: type:name)")
            split_strategy_input = gr.Radio(choices=["token", "recursive"], label="Split Strategy", value="recursive")
            chunk_size_input = gr.Slider(100, 1000, step=100, value=500, label="Chunk Size")
            overlap_size_input = gr.Slider(0, 100, step=10, value=50, label="Overlap Size")
            custom_separators_input = gr.Textbox(label="Custom Split Separators (comma-separated, optional)")
            vector_store_type_input = gr.Radio(choices=["FAISS", "Chroma"], label="Vector Store Type", value="FAISS")
            search_type_input = gr.Radio(choices=["similarity", "mmr", "custom"], label="Search Type", value="similarity")
            lang_input = gr.Dropdown(choices=["german", "english", "french"], label="Language", value="german")

        with gr.Tab("Expert"):
            apply_preprocessing_input = gr.Checkbox(label="Apply Text Preprocessing", value=False)
            optimize_vocab_input = gr.Checkbox(label="Optimize Vocabulary", value=False)
            apply_phonetic_input = gr.Checkbox(label="Apply Phonetic Matching", value=False)
            phonetic_weight_input = gr.Slider(0, 1, step=0.1, value=0.3, label="Phonetic Matching Weight")
            custom_tokenizer_file_input = gr.File(label="Custom Tokenizer File (Optional)")
            custom_tokenizer_model_input = gr.Textbox(label="Custom Tokenizer Model (e.g., WordLevel, BPE, Unigram)")
            custom_tokenizer_vocab_size_input = gr.Textbox(label="Custom Tokenizer Vocab Size", value="10000")
            custom_tokenizer_special_tokens_input = gr.Textbox(label="Custom Tokenizer Special Tokens (comma-separated)")
            use_query_optimization_input = gr.Checkbox(label="Use Query Optimization", value=False)
            query_optimization_model_input = gr.Textbox(label="Query Optimization Model (google/flan-t5-base) ", value="")
            use_reranking_input = gr.Checkbox(label="Use Reranking", value=False)

        with gr.Tab("Automation"):
            with gr.Row():
                auto_file_input = gr.File(label="Upload File (Optional)")
                auto_query_input = gr.Textbox(label="Search Query")

            with gr.Row():
                auto_expected_result_input = gr.Textbox(
                    label="Expected Result (Optional)",
                    placeholder="Enter expected text if you want to evaluate accuracy"
                )
                model_feedback_input = gr.Textbox(
                    label="Model Feedback (Optional)",
                    placeholder="Enter any feedback about model performance"
                )

            with gr.Row():
                with gr.Column():
                    # Default model selection
                    default_models_input = gr.CheckboxGroup(
                        choices=[f"{type}:{name}"
                                for type, names in DEFAULT_MODELS.items()
                                for name in names],
                        label="Default Models",
                        value=[f"HuggingFace:{DEFAULT_MODELS['HuggingFace'][0]}"]
                    )

                with gr.Column():
                    # Custom model input
                    custom_models_input = gr.TextArea(
                        label="Custom Models (Optional)",
                        placeholder="Enter one model per line in format: type:name",
                        lines=3
                    )

            auto_split_strategies = gr.CheckboxGroup(
                choices=["token", "recursive"],
                label="Split Strategies to Test"
            )
            auto_chunk_sizes = gr.TextArea(label="Chunk Sizes to Test (comma-separated)")
            auto_overlap_sizes = gr.TextArea(label="Overlap Sizes to Test (comma-separated)")
            auto_vector_store_types = gr.CheckboxGroup(
                choices=["FAISS", "Chroma"],
                label="Vector Store Types to Test"
            )
            auto_search_types = gr.CheckboxGroup(
                choices=["similarity", "mmr", "custom"],
                label="Search Types to Test"
            )
            auto_top_k = gr.TextArea(label="Top K Values to Test (comma-separated)")
            auto_optimize_vocab = gr.Checkbox(label="Test Vocabulary Optimization", value=True)
            auto_use_query_optimization = gr.Checkbox(label="Test Query Optimization", value=True)
            auto_use_reranking = gr.Checkbox(label="Test Reranking", value=True)

            auto_results_output = gr.Dataframe(label="Automated Test Results", interactive=False)
            auto_stats_output = gr.Dataframe(label="Automated Test Statistics", interactive=False)
            recommendations_output = gr.JSON(label="Recommendations")

            def run_automation(file_input, query_input, expected_result, default_models, custom_models,
                              split_strategies, chunk_sizes, overlap_sizes,
                              vector_store_types, search_types, top_k_values,
                              optimize_vocab, use_query_optimization, use_reranking,
                              model_feedback):
                """Wrapper function to handle Gradio inputs and run automated tests"""

                # Parse model configurations
                model_configs = parse_model_selections(default_models, custom_models)

                # Parse test parameters
                test_params = {
                    'split_strategy': split_strategies,
                    'chunk_size': parse_comma_separated(chunk_sizes),
                    'overlap_size': parse_comma_separated(overlap_sizes),
                    'vector_store_type': vector_store_types,
                    'search_type': search_types,
                    'top_k': parse_comma_separated(top_k_values),
                    'optimize_vocab': [optimize_vocab],
                    'use_query_optimization': [use_query_optimization],
                    'use_reranking': [use_reranking],
                    'lang': ['en'],  # Default to English
                    'apply_preprocessing': [True],  # Default preprocessing
                    'apply_phonetic': [False],  # Default phonetic settings
                    'phonetic_weight': [0.5],
                    'custom_separators': [None],
                    'query_optimization_model': ['google/flan-t5-base']  # Default query optimization model
                }

                # Run automated tests
                results_df, stats_df = run_automated_tests(
                    file_input.name if file_input else None,
                    query_input,
                    model_configs,
                    test_params,
                    expected_result if expected_result else None,
                    model_feedback if model_feedback else None
                )

                # Generate recommendations based on results
                recommendations =  analyze_results(stats_df)

                return results_df, stats_df, recommendations

            auto_submit_button = gr.Button("Run Automated Tests")
            auto_submit_button.click(
                fn=run_automation,
                inputs=[
                    auto_file_input, auto_query_input, auto_expected_result_input,
                    default_models_input, custom_models_input,
                    auto_split_strategies, auto_chunk_sizes, auto_overlap_sizes,
                    auto_vector_store_types, auto_search_types, auto_top_k,
                    auto_optimize_vocab, auto_use_query_optimization, auto_use_reranking,
                    model_feedback_input
                ],
                outputs=[auto_results_output, auto_stats_output, recommendations_output]
            )
        ###
 

        with gr.Tab("LLM Suggestions"):
            llm_file_input = gr.File(label="Upload File for LLM Suggestions")
            llm_num_chunks = gr.Slider(1, 10, step=1, value=5, label="Number of Sample Chunks")
            llm_suggest_button = gr.Button("Get LLM Suggestions")
            llm_suggestions_output = gr.JSON(label="LLM-suggested Settings")

        llm_suggest_button.click(
            fn=get_llm_suggested_settings,
            inputs=[llm_file_input, llm_num_chunks],
            outputs=[llm_suggestions_output]
        ).then(
            fn=update_inputs_with_llm_suggestions,
            inputs=[llm_suggestions_output],
            outputs=[
                embedding_models_input, split_strategy_input, chunk_size_input,
                overlap_size_input, vector_store_type_input, search_type_input,
                top_k_input, apply_preprocessing_input, optimize_vocab_input,
                apply_phonetic_input, phonetic_weight_input
            ]
        )

        results_output = gr.Dataframe(label="Results", interactive=False)
        stats_output = gr.Dataframe(label="Statistics", interactive=False)
        plot_output = gr.Plot(label="Visualizations")
        best_settings_output = gr.JSON(label="Best Settings")

        submit_button = gr.Button("Compare Embeddings")
        submit_button.click(
            #fn=lambda *args: compare_and_show_best(*args),
            fn=lambda *args: compare_embeddings(*args),
            inputs=[
                file_input, query_input, embedding_models_input, custom_embedding_model_input,
                split_strategy_input, chunk_size_input, overlap_size_input, custom_separators_input,
                vector_store_type_input, search_type_input, top_k_input, expected_result_input, lang_input,
                apply_preprocessing_input, optimize_vocab_input, apply_phonetic_input,
                phonetic_weight_input, custom_tokenizer_file_input, custom_tokenizer_model_input,
                custom_tokenizer_vocab_size_input, custom_tokenizer_special_tokens_input,
                use_query_optimization_input, query_optimization_model_input, use_reranking_input
            ],
            outputs=[results_output, stats_output, plot_output, best_settings_output]
        )



    use_case_md = """
    # 🚀 AI Act Embedding Use Case Guide

## 📚 Use Case: Embedding the German AI Act for Local Chat Retrieval

In this guide, we'll walk through the process of embedding the German version of the AI Act using our advanced embedding tool and MTEB. We'll then use these embeddings in a local chat application as a retriever/context.

### Step 1: Prepare the Document 📄

1. Download the German version of the AI Act (let's call it `ai_act_de.txt`).
2. Place the file in your project directory.

### Step 2: Set Up the Embedding Tool 🛠️

1. Open the Embedding Comparison Tool.
2. Navigate to the new "Automation" tab.

### Step 3: Configure the Automated Test 🔧

In the "Use Case" tab, set up the following configuration:

```markdown
- File: ai_act_de.txt
- Query: "Wie definiert das Gesetz KI-Systeme?"
- Model Types: ✅ HuggingFace, ✅ Sentence Transformers
- Model Names: paraphrase-multilingual-MiniLM-L12-v2, distiluse-base-multilingual-cased-v2
- Split Strategies: ✅ recursive, ✅ token
- Chunk Sizes: 256, 512, 1024
- Overlap Sizes: 32, 64, 128
- Vector Store Types: ✅ FAISS
- Search Types: ✅ similarity, ✅ mmr
- Top K Values: 3, 5, 7
- Test Vocabulary Optimization: ✅
- Test Query Optimization: ✅
- Test Reranking: ✅
```

### Step 4: Run the Automated Test 🏃‍♂️

Click the "Run Automated Tests" button and wait for the results.

### Step 5: Analyze the Results 📊

Let's say we got the following simulated results:

```markdown
Best Model: Sentence Transformers - paraphrase-multilingual-MiniLM-L12-v2
Best Settings:
- Split Strategy: recursive
- Chunk Size: 512
- Overlap Size: 64
- Vector Store Type: FAISS
- Search Type: mmr
- Top K: 5
- Optimize Vocabulary: True
- Use Query Optimization: True
- Use Reranking: True

Performance Summary:
- Search Time: 0.15s
- Result Diversity: 0.82
- Rank Correlation: 0.91
- Silhouette Score: 0.76
```

### Step 6: Understand the Results 🧠

1. **Model**: The Sentence Transformers model performed better, likely due to its multilingual capabilities and fine-tuning for paraphrasing tasks.

2. **Split Strategy**: Recursive splitting worked best, probably because it respects the document's structure better than fixed-length token splitting.

3. **Chunk Size**: 512 tokens provide a good balance between context and specificity.

4. **Search Type**: MMR (Maximum Marginal Relevance) outperformed simple similarity search, likely due to its ability to balance relevance and diversity in results.

5. **Optimizations**: All optimizations (vocabulary, query, and reranking) proved beneficial, indicating that the extra processing time is worth the improved results.

### Step 7: Implement in Local Chat 💬

Now that we have the optimal settings, let's implement this in a local chat application:

1. Use the `paraphrase-multilingual-MiniLM-L12-v2` model for embeddings.
2. Set up a FAISS vector store with the embedded chunks.
3. Implement MMR search with a top-k of 5.
4. Include the optimization steps in your pipeline.

### Step 8: Test the Implementation 🧪

Create a simple chat interface and test with various queries about the AI Act. For example:

User: "Was sind die Hauptziele des KI-Gesetzes?"
    """


    tutorial_md = """
# Advanced Embedding Comparison Tool Tutorial

Welcome to the **Advanced Embedding Comparison Tool**! This comprehensive guide will help you understand and utilize the tool's features to optimize your **Retrieval-Augmented Generation (RAG)** systems.

## Table of Contents
1. [Introduction to RAG](#introduction-to-rag)
2. [Key Components of RAG](#key-components-of-rag)
3. [Impact of Parameter Changes](#impact-of-parameter-changes)
4. [Advanced Features](#advanced-features)
5. [Using the Embedding Comparison Tool](#using-the-embedding-comparison-tool)
6. [Automated Testing and Analysis](#automated-testing-and-analysis)
7. [Mathematical Concepts and Metrics](#mathematical-concepts-and-metrics)
8. [Code Examples](#code-examples)
9. [Best Practices and Tips](#best-practices-and-tips)
10. [Resources and Further Reading](#resources-and-further-reading)

---

## Introduction to RAG

**Retrieval-Augmented Generation (RAG)** is a powerful technique that combines the strengths of large language models (LLMs) with the ability to access and use external knowledge. RAG is particularly useful for:

- Providing up-to-date information
- Answering questions based on specific documents or data sources
- Reducing hallucinations in AI responses
- Customizing AI outputs for specific domains or use cases

RAG is ideal for applications requiring accurate, context-specific information retrieval combined with natural language generation, such as chatbots, question-answering systems, and document analysis tools.

---

## Key Components of RAG

### 1. Document Loading
Ingests documents from various sources (PDFs, web pages, databases, etc.) into a format that can be processed by the RAG system. The tool supports multiple file formats, including PDF, DOCX, and TXT.

### 2. Document Splitting
Splits large documents into smaller chunks for more efficient processing and retrieval. Available strategies include:
- **Token-based splitting**
- **Recursive splitting**

### 3. Vector Store and Embeddings
Embeddings are dense vector representations of text that capture semantic meaning. The tool supports multiple embedding models and vector stores:
- **Embedding models**: HuggingFace, OpenAI, Cohere, and custom models.
- **Vector stores**: FAISS and Chroma.

### 4. Retrieval
Finds the most relevant documents or chunks based on a query. Available retrieval methods include:
- **Similarity search**
- **Maximum Marginal Relevance (MMR)**
- **Custom search methods**

---

## Impact of Parameter Changes

Understanding how different parameters affect your RAG system is crucial for optimization:

- **Chunk Size**: Larger chunks provide more context but may reduce precision. Smaller chunks increase precision but may lose context.
- **Overlap**: More overlap helps maintain context between chunks but increases computational load.
- **Embedding Model**: Performance varies across languages and domains.
- **Vector Store**: Affects query speed and the types of searches.
- **Retrieval Method**: Influences the diversity and relevance of retrieved documents.

---

## Advanced Features

### 1. Custom Tokenization
Upload a custom tokenizer file and specify the tokenizer model, vocabulary size, and special tokens for domain or language-specific tokenization.

### 2. Query Optimization
Improve search results by generating multiple variations of the input query using a language model to capture different phrasings.

### 3. Reranking
Further refine search results by using a separate model to re-score and reorder the initial retrieval results.

### 4. Phonetic Matching
For languages like German, phonetic matching with adjustable weighting is available.

### 5. Vocabulary Optimization
Optimize vocabulary for domain-specific applications during the embedding process.

---

## Using the Embedding Comparison Tool

The tool is divided into several tabs for ease of use:

### Simple Tab
1. **File Upload**: Upload a file (PDF, DOCX, or TXT) or use files from the `./files` directory.
2. **Search Query**: Enter the search query.
3. **Embedding Models**: Select one or more embedding models to compare.
4. **Top K**: Set the number of top results to retrieve (1-10).

### Advanced Tab
5. **Custom Embedding Model**: Specify a custom embedding model.
6. **Split Strategy**: Choose between 'token' and 'recursive' splitting.
7. **Chunk Size**: Set chunk size (100-1000).
8. **Overlap Size**: Set overlap between chunks (0-100).
9. **Custom Split Separators**: Enter custom separators for text splitting.
10. **Vector Store Type**: Choose between FAISS and Chroma.
11. **Search Type**: Select 'similarity', 'mmr', or 'custom'.
12. **Language**: Specify the document's primary language.

### Optional Tab
13. **Text Preprocessing**: Toggle text preprocessing.
14. **Vocabulary Optimization**: Enable vocabulary optimization.
15. **Phonetic Matching**: Enable phonetic matching and set its weight.
16. **Custom Tokenizer**: Upload a custom tokenizer and specify parameters.
17. **Query Optimization**: Enable query optimization and specify the model.
18. **Reranking**: Enable result reranking.

---

## Automated Testing and Analysis

The **Automation tab** allows you to run comprehensive tests across multiple configurations:

1. Set up test parameters like model types, split strategies, chunk sizes, etc.
2. Click "Run Automated Tests."
3. View results, statistics, and recommendations to find optimal configurations for your use case.

---

## Mathematical Concepts and Metrics

### Cosine Similarity
Measures the cosine of the angle between two vectors, used in similarity search:
$$\text{cosine similarity} = \frac{\mathbf{A} \cdot \mathbf{B}}{\|\mathbf{A}\| \|\mathbf{B}\|}$$

### Maximum Marginal Relevance (MMR)
Balances relevance and diversity in search results:
$$\text{MMR} = \arg\max_{D_i \in R \setminus S} [\lambda \text{Sim}_1(D_i, Q) - (1-\lambda) \max_{D_j \in S} \text{Sim}_2(D_i, D_j)]$$

### Silhouette Score
Measures how well an object fits within its own cluster compared to others. Scores range from -1 to 1, where higher values indicate better-defined clusters.

---

## Code Examples

### Custom Tokenization
```python
def create_custom_tokenizer(file_path, model_type='WordLevel', vocab_size=10000, special_tokens=None):
    with open(file_path, 'r', encoding='utf-8') as f:
        text = f.read()

    tokenizer = Tokenizer(models.WordLevel(unk_token="[UNK]")) if model_type == 'WordLevel' else Tokenizer(models.BPE(unk_token="[UNK]"))
    tokenizer.pre_tokenizer = Whitespace()

    trainer = trainers.WordLevelTrainer(special_tokens=special_tokens or ["[UNK]", "[CLS]", "[SEP]", "[PAD]", "[MASK]"], vocab_size=vocab_size)
    tokenizer.train_from_iterator([text], trainer)

    return tokenizer
````

### Query Optimization
```python
def optimize_query(query, llm):
    multi_query_retriever = MultiQueryRetriever.from_llm(
        retriever=get_retriever(vector_store, search_type, search_kwargs),
        llm=llm
    )
    optimized_queries = multi_query_retriever.invoke(query)
    return optimized_queries
````

### Reranking
```python
def rerank_results(results, query, reranker):
    reranked_results = reranker.rerank(query, [doc.page_content for doc in results])
    return reranked_results
````

### Best Practices and Tips

- Start Simple: Begin with basic configurations, then gradually add complexity.
- Benchmark: Use automated testing to benchmark different setups.
- Domain-Specific Tuning: Consider custom tokenizers and embeddings for specialized domains.
- Balance Performance and Cost: Use advanced features like query optimization and reranking judiciously.
- Iterate: Optimization is an iterative process—refine your approach based on tool insights.


    ## Useful Resources and Links

    Here are some valuable resources to help you better understand and work with embeddings, retrieval systems, and natural language processing:

    ### Embeddings and Vector Databases
    - [Understanding Embeddings](https://www.tensorflow.org/text/guide/word_embeddings): A guide by TensorFlow on word embeddings
    - [FAISS: A Library for Efficient Similarity Search](https://github.com/facebookresearch/faiss): Facebook AI's vector similarity search library
    - [Chroma: The AI-native open-source embedding database](https://www.trychroma.com/): An embedding database designed for AI applications

    ### Natural Language Processing
    - [NLTK (Natural Language Toolkit)](https://www.nltk.org/): A leading platform for building Python programs to work with human language data
    - [spaCy](https://spacy.io/): Industrial-strength Natural Language Processing in Python
    - [Hugging Face Transformers](https://huggingface.co/transformers/): State-of-the-art Natural Language Processing for PyTorch and TensorFlow 2.0

    ### Retrieval-Augmented Generation (RAG)
    - [LangChain](https://python.langchain.com/docs/get_started/introduction): A framework for developing applications powered by language models
    - [OpenAI's RAG Tutorial](https://platform.openai.com/docs/tutorials/web-qa-embeddings): A guide on building a QA system with embeddings

    ### German Language Processing
    - [Kölner Phonetik](https://en.wikipedia.org/wiki/Cologne_phonetics): Information about the Kölner Phonetik algorithm
    - [German NLP Resources](https://github.com/adbar/German-NLP): A curated list of open-access resources for German NLP

    ### Benchmarks and Evaluation
    - [MTEB Leaderboard](https://huggingface.co/spaces/mteb/leaderboard): Massive Text Embedding Benchmark leaderboard
    - [GLUE Benchmark](https://gluebenchmark.com/): General Language Understanding Evaluation benchmark

    ### Tools and Libraries
    - [Gensim](https://radimrehurek.com/gensim/): Topic modelling for humans
    - [Sentence-Transformers](https://www.sbert.net/): A Python framework for state-of-the-art sentence, text and image embeddings

       ### Support me
    - [Visual Crew Builder](https://visual-crew.builder.ai/): Tool for create AI systems, workflows and api. Or just a notebook.



This tool empowers you to fine-tune your RAG system for optimal performance. Experiment with different settings, run automated tests, and use insights to create an efficient information retrieval and generation system.

# Template

python
´´´
# Chat App Template
def create_chat_app(settings):
    def chat(message, history):
        # Process the message using the configured embedding model and vector store
        chunks, embedding_model, _ = process_files(
            settings['file_path'],
            settings['model_type'],
            settings['model_name'],
            settings['split_strategy'],
            settings['chunk_size'],
            settings['overlap_size'],
            settings['custom_separators'],
            settings['lang'],
            settings['apply_preprocessing']
        )

        results, _, _, _ = search_embeddings(
            chunks,
            embedding_model,
            settings['vector_store_type'],
            settings['search_type'],
            message,
            settings['top_k'],
            lang=settings['lang'],
            apply_phonetic=settings['apply_phonetic'],
            phonetic_weight=settings['phonetic_weight']
        )

        # Generate a response based on the retrieved results
        response = f"Based on the query '{message}', here are the top {settings['top_k']} relevant results:\n\n"
        for i, result in enumerate(results[:settings['top_k']]):
            response += f"{i+1}. {result['content'][:100]}...\n\n"

        return response

    with gr.Blocks() as chat_interface:
        gr.Markdown(f"# Chat App using {settings['model_type']} - {settings['model_name']}")
        chatbot = gr.Chatbot()
        msg = gr.Textbox()
        clear = gr.Button("Clear")

        msg.submit(chat, [msg, chatbot], [msg, chatbot])
        clear.click(lambda: None, None, chatbot, queue=False)

    return chat_interface

# Sample usage of the chat app template
sample_settings = {
    'file_path': 'path/to/your/document.pdf',
    'model_type': 'HuggingFace',
    'model_name': 'paraphrase-miniLM',
    'split_strategy': 'recursive',
    'chunk_size': 500,
    'overlap_size': 50,
    'custom_separators': None,
    'vector_store_type': 'FAISS',
    'search_type': 'similarity',
    'top_k': 3,
    'lang': 'english',
    'apply_preprocessing': True,
    'apply_phonetic': True,
    'phonetic_weight': 0.3
}

sample_chat_app = create_chat_app(sample_settings)

if __name__ == "__main__":
    launch_interface()
    # Uncomment the following line to launch the sample chat app
´´´

        """


    iface = gr.TabbedInterface(
        [iface, gr.Markdown(tutorial_md), gr.Markdown( use_case_md )],
        ["Embedding Comparison", "Tutorial", "Use Case"]
    )

    iface.launch(debug=True, share=True)

# Enhanced Automated Testing
def run_automated_tests(file_path: str, query: str, model_configs: List[Dict[str, str]],
                       test_params: Dict[str, List[Any]], expected_result: Optional[str] = None,
                       model_feedback: Optional[str] = None) -> Tuple[pd.DataFrame, pd.DataFrame]:
    """
    Enhanced automated testing function with support for custom models and feedback
    """
    all_results = []
    all_stats = []
    model_manager = ModelManager()

    # Create parameter grid excluding model configurations
    base_params = {k: v for k, v in test_params.items() if k not in ['model_type', 'model_name']}
    param_grid = ParameterGrid(base_params)

    # Test each model configuration with all parameter combinations
    for model_config in tqdm(model_configs, desc="Testing models"):
        model_type = model_config['type']
        model_name = model_config['name']

        for params in tqdm(param_grid, desc=f"Testing parameters for {model_type}:{model_name}"):
            try:
                # Process files and get chunks
                chunks, embedding_model, num_tokens = process_files(
                    file_path,
                    model_type,
                    model_name,
                    params['split_strategy'],
                    params['chunk_size'],
                    params['overlap_size'],
                    params.get('custom_separators'),
                    params['lang'],
                    params['apply_preprocessing']
                )

                # Apply vocabulary optimization if specified
                if params['optimize_vocab']:
                    tokenizer, chunks = optimize_vocabulary(chunks)

                # Apply query optimization if specified
                current_query = query
                if params['use_query_optimization']:
                    optimized_queries = optimize_query(
                        query,
                        params['query_optimization_model'],
                        chunks,
                        embedding_model,
                        params['vector_store_type'],
                        params['search_type'],
                        params['top_k']
                    )
                    current_query = " ".join(optimized_queries)

                # Perform search
                results, search_time, vector_store, raw_results = search_embeddings(
                    chunks,
                    embedding_model,
                    params['vector_store_type'],
                    params['search_type'],
                    current_query,
                    params['top_k'],
                    expected_result,
                    params['lang'],
                    params['apply_phonetic'],
                    params['phonetic_weight']
                )

                # Apply reranking if specified
                if params['use_reranking']:
                    reranker = pipeline("text-classification",
                                      model="cross-encoder/ms-marco-MiniLM-L-12-v2")
                    raw_results = rerank_results(raw_results, current_query, reranker)

                # Calculate statistics
                stats = ResultAnalyzer.calculate_statistics(
                    raw_results, search_time, vector_store, num_tokens,
                    embedding_model, current_query, params['top_k'],
                    expected_result, model_feedback
                )

                # Update model rankings
                model_id = f"{model_type}:{model_name}"
                ranking_score = calculate_model_ranking_score(stats)
                model_manager.update_model_ranking(model_id, ranking_score, model_feedback)

                # Add model information to stats
                stats.update({
                    "model_type": model_type,
                    "model_name": model_name,
                    "model": f"{model_type} - {model_name}",
                    **params
                })

                # Format and store results
                all_results.extend(format_results(raw_results, stats))
                all_stats.append(stats)

            except Exception as e:
                print(f"Error testing {model_type}:{model_name} with parameters {params}: {str(e)}")
                continue

    return pd.DataFrame(all_results), pd.DataFrame(all_stats)

    # Helper function to calculate model ranking score
def calculate_model_ranking_score(stats: Dict[str, Any]) -> float:
    """Calculate a composite score for model ranking"""
    weights = {
        'search_time': -0.2,  # Negative weight because lower is better
        'result_diversity': 0.2,
        'rank_correlation': 0.3,
        'contains_expected': 0.3,
        'expected_result_rank': -0.2  # Negative weight because lower rank is better
    }

    score = 0.0
    for metric, weight in weights.items():
        if metric in stats and not isinstance(stats[metric], str):
            if metric == 'contains_expected':
                value = float(stats[metric])
            elif metric == 'expected_result_rank':
                value = 1.0 / max(stats[metric], 1)  # Convert rank to score (higher is better)
            else:
                value = float(stats[metric])
            score += weight * value

    return score

if __name__ == "__main__":
    launch_interface()
